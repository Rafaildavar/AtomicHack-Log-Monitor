"""Сервис для генерации отчетов анализа логов."""

import json
import os
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.shared import Inches
from aiogram.types import InputFile


class ReportGenerator:
    """Генератор отчетов в различных форматах."""

    def __init__(self, output_dir: str = "reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    async def generate_docx_report(self, analysis: Dict, file_name: str, user_name: str = "Пользователь") -> str:
        """Генерирует отчет в формате DOCX."""
        doc = Document()

        # Заголовок
        title = doc.add_heading('Отчет анализа логов', 0)
        title.alignment = 1  # Центрирование

        # Информация об анализе
        doc.add_heading('Информация об анализе', level=1)
        doc.add_paragraph(f'Файл: {file_name}')
        doc.add_paragraph(f'Пользователь: {user_name}')
        doc.add_paragraph(f'Дата анализа: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        # Статистика
        doc.add_heading('Статистика', level=1)
        doc.add_paragraph(f'Общее количество строк: {analysis["total_lines"]}')
        doc.add_paragraph(f'Ошибок (ERROR/CRITICAL/FATAL): {analysis["error_count"]}')
        doc.add_paragraph(f'Предупреждений (WARNING): {analysis["warning_count"]}')
        doc.add_paragraph(f'Информационных сообщений (INFO): {analysis["info_count"]}')

        # Распределение по уровням
        if analysis.get("level_distribution"):
            doc.add_heading('Распределение по уровням', level=2)
            for level, count in analysis["level_distribution"].items():
                doc.add_paragraph(f'• {level}: {count}')

        # Временной диапазон
        if analysis["time_range"]:
            doc.add_heading('Временной диапазон', level=2)
            doc.add_paragraph(f'Начало: {analysis["time_range"]["start"]}')
            doc.add_paragraph(f'Конец: {analysis["time_range"]["end"]}')

        # Источники
        if analysis["sources"]:
            doc.add_heading('Источники', level=2)
            for source, count in analysis["sources"].items():
                doc.add_paragraph(f'• {source}: {count}')

        # Наиболее частые сообщения
        if analysis.get("top_messages"):
            doc.add_heading('Наиболее частые сообщения', level=2)
            for msg, count in analysis["top_messages"].items():
                doc.add_paragraph(f'• {msg}: {count} раз')

        # Детальные данные (если есть)
        if analysis.get("dataframe"):
            doc.add_heading('Детальные данные', level=2)
            doc.add_paragraph('Таблица с детальными записями логов прилагается в формате JSON.')

        # Сохраняем файл
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_filename = f"log_analysis_{timestamp}.docx"
        report_path = os.path.join(self.output_dir, report_filename)

        doc.save(report_path)
        return report_path

    async def generate_json_report(self, analysis: Dict, file_name: str) -> str:
        """Генерирует отчет в формате JSON."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_filename = f"log_analysis_{timestamp}.json"
        report_path = os.path.join(self.output_dir, report_filename)

        # Подготавливаем данные для JSON
        report_data = {
            "metadata": {
                "file_name": file_name,
                "analysis_date": datetime.now().isoformat(),
                "format_version": "1.0"
            },
            "statistics": {
                "total_lines": analysis["total_lines"],
                "error_count": analysis["error_count"],
                "warning_count": analysis["warning_count"],
                "info_count": analysis["info_count"],
                "time_range": analysis["time_range"],
                "sources": analysis["sources"],
                "top_messages": analysis["top_messages"],
                "level_distribution": analysis["level_distribution"]
            }
        }

        # Добавляем детальные данные если есть
        if analysis.get("dataframe"):
            report_data["detailed_data"] = analysis["dataframe"]

        # Сохраняем JSON
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False)

        return report_path

    async def get_input_file(self, file_path: str) -> InputFile:
        """Возвращает InputFile для отправки через Telegram."""
        return InputFile(file_path, filename=os.path.basename(file_path))
